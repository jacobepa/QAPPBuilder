# test_export_docx_views.py
# !/usr/bin/env python3
# coding=utf-8

from django.test import TestCase, Client
from docx import Document
from io import BytesIO
import constants.qapp_section_a_const as constants_a
from qapp_builder.tests.mixins import QappTestMixin


class TestExportDocxViews(QappTestMixin, TestCase):
    """Tests for the export_docx_views module."""

    def setUp(self):
        """Set up test data."""
        # Create test user and QAPP
        self.user = self.create_test_user()
        self.qapp = self.create_test_qapp(self.user)

        # Create all test sections
        sections = self.create_all_test_sections(self.qapp)
        self.section_a1 = sections['section_a1']
        self.section_a2 = sections['section_a2']
        self.section_a4 = sections['section_a4']
        self.section_a5 = sections['section_a5']
        self.section_a6 = sections['section_a6']
        self.section_a10 = sections['section_a10']
        self.section_a11 = sections['section_a11']
        self.section_b = sections['section_b']
        self.section_b7 = sections['section_b7']
        self.section_c = sections['section_c']
        self.section_d = sections['section_d']

        # Set up client
        self.client = Client()
        self.client.login(username='testuser', password='12345')

    def test_export_qapp_docx_authenticated(self):
        """Test exporting QAPP as DOCX when authenticated."""
        response = self.client.get(f'/qapp/{self.qapp.id}/export/?format=docx')

        # Check response status code
        self.assertEqual(response.status_code, 200)

        # Check content type
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        self.assertEqual(response['Content-Type'], content_type)

        # Check filename
        expected_filename = f'{self.qapp.title} - {constants_a.QAPP_STR}.docx'
        self.assertEqual(
            response['Content-Disposition'],
            f'attachment; filename="{expected_filename}"'
        )

        # Verify document content
        doc = Document(BytesIO(response.content))

        # Check document has content
        self.assertTrue(len(doc.paragraphs) > 0)

        # Check for section headers
        section_headers = []
        for paragraph in doc.paragraphs:
            # Check if paragraph is a heading
            if paragraph.style.name.startswith('Heading'):
                section_headers.append(paragraph.text)

        # Verify we found at least one section header
        self.assertTrue(len(section_headers) > 0, "No section headers found in document")

        # Check for specific section headers
        expected_headers = [
            'A1: Title Page',
            'A2: Approval Page',
            'A3: Table of Contents'
        ]

        found_headers = [
            header for header in expected_headers
            if any(header in h for h in section_headers)
        ]

        self.assertTrue(
            len(found_headers) > 0,
            f"No expected section headers found. Expected one of: {expected_headers}"
        )

    def test_export_qapp_docx_unauthenticated(self):
        """Test exporting QAPP as DOCX when not authenticated."""
        # Logout the client
        self.client.logout()

        response = self.client.get(f'/qapp/{self.qapp.id}/export/?format=docx')

        # Should redirect to login page
        self.assertEqual(response.status_code, 302)
        self.assertTrue('/accounts/login/' in response.url)

    def test_export_qapp_docx_nonexistent(self):
        """Test exporting non-existent QAPP."""
        response = self.client.get('/qapp/-1/export/?format=docx')
        # Should return 404
        self.assertEqual(response.status_code, 404)

    def test_export_qapp_docx_content(self):
        """Test the content of the exported DOCX file."""
        response = self.client.get(f'/qapp/{self.qapp.id}/export/?format=docx')

        doc = Document(BytesIO(response.content))

        # Check for QAPP title
        title_found = False
        for paragraph in doc.paragraphs:
            if self.qapp.title in paragraph.text:
                title_found = True
                break
        self.assertTrue(title_found)

        # Check for Section A1 content
        section_a1_content = [
            self.section_a1.ord_center,
            self.section_a1.division,
            self.section_a1.branch,
            self.section_a1.ord_national_program,
            self.section_a1.proj_qapp_id,
            self.section_a1.qa_category
        ]

        content_found = {content: False for content in section_a1_content}
        for paragraph in doc.paragraphs:
            for content in section_a1_content:
                if content in paragraph.text:
                    content_found[content] = True

        # All Section A1 content should be found
        self.assertTrue(all(content_found.values()))