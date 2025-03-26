from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def parse_fy(fy):
    # Remove any non-numeric characters
    fy = ''.join(filter(str.isdigit, str(fy)))
    # Return the last two digits
    return int(fy[-2:])


def set_font(run, name='Calibri (Body)', size=11, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_fake_header_style(paragraph):
    font_name = 'Calibri Light (Heading)'
    font_size = 16
    bold = True
    color = '2F5496'

    run = paragraph.add_run()
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)

    # Set paragraph properties
    p = paragraph._element
    pPr = p.get_or_add_pPr()

    # Create and append rPr element
    rPr = OxmlElement('w:rPr')

    # Set custom style properties
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))  # font size in half-points
    rPr.append(sz)

    b = OxmlElement('w:b')
    b.set(qn('w:val'), 'true' if bold else 'false')
    rPr.append(b)

    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), color)
    rPr.append(color_elem)

    pPr.append(rPr)


def create_toc(doc, c, format_type):
    """Set up the Table of Contents page."""
    if format_type == 'docx':
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        # creates a new element
        field_char = OxmlElement('w:fldChar')
        # sets attribute on element
        field_char.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        # sets attribute on element
        instr_text.set(qn('xml:space'), 'preserve')
        # change 1-3 depending on heading levels you need
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

        field_char2 = OxmlElement('w:fldChar')
        field_char2.set(qn('w:fldCharType'), 'separate')
        field_char3 = OxmlElement('w:t')
        field_char3.text = "Right-click to update Table of Contents."
        field_char2.append(field_char3)

        field_char4 = OxmlElement('w:fldChar')
        field_char4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(field_char)
        r_element.append(instr_text)
        r_element.append(field_char2)
        r_element.append(field_char4)
    elif format_type == 'pdf':
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, c._height - 100, "Table of Contents")
        c._height -= 40  # Adjust height for next element

        toc_entries = [
            "1. Section A1 - Introduction",
            "2. Section A2 - QAPP Approvals",
            "3. Section A3 - Revision History",
            "4. Section A4 - Project Background",
            "5. Section A5 - Project Task Schedule",
            "6. Section A6 - Information",
            "7. Section A7 - Distribution List",
            "8. Section A8 - Roles and Responsibilities",
            "9. Section A9 - Accessibility",
            "10. Section A10 - Organization Chart",
            "11. Section A11 - Additional Information",
            "12. Section A12 - Document Records",
            "13. Section B - Project Description",
            "14. Section C - Quality Assurance",
            "15. Section D - Data Management"
        ]

        c.setFont("Helvetica", 12)
        for entry in toc_entries:
            c.drawString(100, c._height - 20, entry)
            c._height -= 20  # Adjust height for next element

        c.showPage()


def add_heading(doc, c, text, level, format_type):
    if format_type == 'docx':
        doc.add_heading(text, level=level)
    elif format_type == 'pdf':
        font_size = {1: 16, 2: 14, 3: 12}.get(level, 12)
        c.setFont("Helvetica-Bold", font_size)
        c.drawString(100, c._height - 100, text)
        c._height -= 20  # Adjust height for next element


def add_paragraph(doc,
                  c,
                  text,
                  format_type,
                  name='Calibri (Body)',
                  size=11,
                  bold=False,
                  italic=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT):
    if format_type == 'docx':
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, name, size, bold, italic)
        p.alignment = alignment
    elif format_type == 'pdf':
        c.setFont("Helvetica", size)
        c.drawString(100, c._height - 100, text)
        c._height -= 20  # Adjust height for next element


def add_table(doc, c, data, format_type, column_widths=None):
    if format_type == 'docx':
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        if column_widths:
            # NOTE: These widths must be integers, so cast them
            column_widths = [int(x) for x in column_widths]
            for i, width in enumerate(column_widths):
                table.columns[i].width = width
        for row_idx, row_data in enumerate(data):
            row_cells = table.rows[row_idx].cells
            for cell_idx, cell_data in enumerate(row_data):
                row_cells[cell_idx].text = cell_data
    elif format_type == 'pdf':
        from reportlab.platypus import Table, TableStyle
        table = Table(data, colWidths=column_widths)
        table.setStyle(TableStyle([('GRID', (0, 0), (-1, -1), 0.5, 'black')]))
        table.wrapOn(c, c._width, c._height)
        table.drawOn(c, 100, c._height - 100)
        c._height -= 20 * len(data)  # Adjust height for next element
