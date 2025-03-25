from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from os.path import exists
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
import constants.qapp_section_a_const as constants_a
import constants.qapp_section_b_const as constants_b
import constants.qapp_section_c_d_const as constants_c_d
from qapp_builder.models import Qapp, SectionA1, Revision, \
    SectionA2, AdditionalSignature, AcronymAbbreviation, SectionA4, SectionA5, \
    SectionA6, Distribution, RoleResponsibility, SectionA10, SectionA11, \
    DocumentRecord, SectionB, SectionB7, HardwareSoftware, SectionC, SectionD
from qapp_builder.views.export_utilities import parse_fy, add_table, \
    set_fake_header_style, create_toc, add_heading, add_paragraph


@login_required
def export_qapp_docx(request, qapp_id):
    qapp = Qapp.objects.get(id=qapp_id)
    file_name = f'{qapp.title} - {constants_a.QAPP_STR}.docx'

    # Create the empty document
    doc = Document()
    c = None  # No canvas for DOCX format
    format_type = 'docx'

    # Write the Sections A:
    write_section_a1(qapp, doc, c, format_type)
    write_section_a2(qapp, doc, c, format_type)
    write_section_a3(qapp, doc, c, format_type)
    write_section_a4(qapp, doc, c, format_type)
    write_section_a5(qapp, doc, c, format_type)
    write_section_a6(qapp, doc, c, format_type)
    write_section_a7(qapp, doc, c, format_type)
    write_section_a8(qapp, doc, c, format_type)
    write_section_a9(qapp, doc, c, format_type)
    write_section_a10(qapp, doc, c, format_type)
    write_section_a11(qapp, doc, c, format_type)
    write_section_a12(qapp, doc, c, format_type)

    # Write the Sections B, C, D:
    write_section_b(qapp, doc, c, format_type)
    write_section_b7(qapp, doc, c, format_type)
    write_section_c(qapp, doc, c, format_type)
    write_section_d(qapp, doc, c, format_type)

    # Save the document to a BytesIO object
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Create the HTTP response with the appropriate headers
    response = HttpResponse(
        file_stream,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')  # noqa: E501
    response['Content-Disposition'] = f'attachment; filename="{file_name}"'

    return response


@login_required
def export_qapp_pdf(request, qapp_id):
    qapp = Qapp.objects.get(id=qapp_id)
    file_name = f'{qapp.title} - {constants_a.QAPP_STR}.pdf'

    # Create the PDF document
    file_stream = BytesIO()
    c = canvas.Canvas(file_stream, pagesize=letter)
    width, height = letter

    # Write the Sections A:
    write_section_a1(qapp, c, width, height, format_type='pdf')
    write_section_a2(qapp, c, width, height, format_type='pdf')
    write_section_a3(qapp, c, width, height, format_type='pdf')
    write_section_a4(qapp, c, width, height, format_type='pdf')
    write_section_a5(qapp, c, width, height, format_type='pdf')
    write_section_a6(qapp, c, width, height, format_type='pdf')
    write_section_a7(qapp, c, width, height, format_type='pdf')
    write_section_a8(qapp, c, width, height, format_type='pdf')
    write_section_a9(qapp, c, width, height, format_type='pdf')
    write_section_a10(qapp, c, width, height, format_type='pdf')
    write_section_a11(qapp, c, width, height, format_type='pdf')
    write_section_a12(qapp, c, width, height, format_type='pdf')
    write_section_b(qapp, c, width, height, format_type='pdf')
    write_section_b7(qapp, c, width, height, format_type='pdf')
    write_section_c(qapp, c, width, height, format_type='pdf')
    write_section_d(qapp, c, width, height, format_type='pdf')

    # Save the PDF document
    c.save()
    file_stream.seek(0)

    # Create the HTTP response with the appropriate headers
    response = HttpResponse(file_stream, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{file_name}"'

    return response


def write_section_a1(qapp, doc, c, format_type):
    section_a1 = SectionA1.objects.get(qapp_id=qapp.id)

    # Write Heading 1
    add_heading(doc,
                c,
                constants_a.SECTION_A['a']['header'],
                level=1,
                format_type=format_type)

    # Write Heading 2
    add_heading(doc,
                c,
                constants_a.SECTION_A['a1']['header'],
                level=2,
                format_type=format_type)

    # Centered, bold, size 16 section
    add_paragraph(doc,
                  c,
                  constants_a.QAPP_TITLE_HEADER,
                  format_type=format_type,
                  size=16,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
                  c,
                  section_a1.ord_center,
                  format_type=format_type,
                  size=16,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
                  c,
                  section_a1.division,
                  format_type=format_type,
                  size=16,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
                  c,
                  section_a1.branch,
                  format_type=format_type,
                  size=16,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
                  c,
                  f'{qapp.title} {constants_a.QAPP_STR}',
                  format_type=format_type,
                  size=16,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Centered, size 12 section
    add_paragraph(doc,
                  c,
                  f'ORD National Program: {section_a1.ord_national_program}',
                  format_type=format_type,
                  size=12,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
                  c,
                  f'Version Date: {section_a1.version_date}',
                  format_type=format_type,
                  size=12,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
                  c,
                  f'Project QAPP ID: {section_a1.proj_qapp_id}',
                  format_type=format_type,
                  size=12,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
                  c,
                  f'QA Category: {section_a1.qa_category}',
                  format_type=format_type,
                  size=12,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
                  c,
                  f'QAPP Developed: {section_a1.intra_or_extra}',
                  format_type=format_type,
                  size=12,
                  bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if section_a1.intra_or_extra == constants_a.INTRAMURALLY:
        add_paragraph(doc,
                      c,
                      constants_a.INTRAMURAL_TEXT,
                      format_type=format_type,
                      size=12,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_paragraph(doc,
                      c,
                      f'Vehicle #: {section_a1.vehicle_num}',
                      format_type=format_type,
                      size=12,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(doc,
                      c,
                      f'Name of Non-EPA Organization: {section_a1.non_epa_org}',
                      format_type=format_type,
                      size=12,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(
            doc,
            c,
            f'Period of Performance (POP): {section_a1.period_performance}',
            format_type=format_type,
            size=12,
            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(doc,
                      c,
                      constants_a.EXTRAMURAL_TEXT,
                      format_type=format_type,
                      size=12,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(
        doc,
        c,
        'QAPP Accessibility: QAPPs will be made internally accessible via the ORD QAPP intranet site upon final approval unless the following statement is selected.',  # noqa: E501
        format_type=format_type,
        size=12,
        italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        c,
        'I do NOT want this QAPP internally shared and accessible on the ORD intranet site.',  # noqa: E501
        format_type=format_type,
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Insert page break
    if format_type == 'docx':
        doc.add_page_break()
    elif format_type == 'pdf':
        c.showPage()


def write_section_a2(qapp, doc, c, format_type):
    section_a2 = SectionA2.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_a.SECTION_A['a2']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  "QAPP Approvals (Electronic Approval Signatures/Dates)",
                  format_type=format_type,
                  size=12,
                  bold=True)

    sig_line_len = 45
    table_data = [[
        f'{label}: ', '_' * sig_line_len
    ] for attribute, label in constants_a.SECTION_A['a2']['labels'].items()]
    additional_sigs = AdditionalSignature.objects.filter(
        section_a2_id=section_a2.id)
    table_data.extend([[f'{sig.title} ({sig.name}): ', '_' * sig_line_len]
                       for sig in additional_sigs])

    add_table(doc,
              c,
              table_data,
              format_type=format_type,
              column_widths=[2 * inch, 4 * inch])

    if format_type == 'docx':
        doc.add_page_break()
    elif format_type == 'pdf':
        c.showPage()


def write_section_a3(qapp, doc, c, format_type):
    section_a1 = SectionA1.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_a.SECTION_A['a3']['header'],
                level=2,
                format_type=format_type)
    create_toc(doc, c, format_type)
    add_paragraph(doc, c, "Revision History", format_type=format_type)
    set_fake_header_style(doc.paragraphs[-1])

    revisions = Revision.objects.filter(qapp_id=qapp.id)
    table_data = [['Date', 'QAPP ID', 'Author(s)', 'Description of Revision']
                  ] + [[
                      str(revision.date), section_a1.proj_qapp_id,
                      revision.author, revision.description
                  ] for revision in revisions]
    add_table(doc,
              c,
              table_data,
              format_type=format_type,
              column_widths=[0.75 * inch, 0.75 * inch, 2 * inch, 2.5 * inch])

    add_paragraph(doc,
                  c,
                  "Acronyms/Abbreviations/Definitions",
                  format_type=format_type)
    set_fake_header_style(doc.paragraphs[-1])

    acronyms = AcronymAbbreviation.objects.filter(qapp_id=qapp.id)
    table_data = [['Acronym/Abbreviation', 'Definition']] + \
        [[acronym.acronym_abbreviation, acronym.definition]
            for acronym in acronyms]
    add_table(doc,
              c,
              table_data,
              format_type=format_type,
              column_widths=[2.5 * inch, 3.5 * inch])

    if format_type == 'docx':
        doc.add_page_break()
    elif format_type == 'pdf':
        c.showPage()


def write_section_a4(qapp, doc, c, format_type):
    section_a4 = SectionA4.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_a.SECTION_A['a4']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_a.SECTION_A['a4']['boilerplate'],
                  format_type=format_type,
                  size=11)
    add_heading(doc,
                c,
                constants_a.SECTION_A['a4']['labels']['project_background'],
                level=3,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  section_a4.project_background,
                  format_type=format_type,
                  size=11)
    add_heading(doc,
                c,
                constants_a.SECTION_A['a4']['labels']['project_purpose'],
                level=3,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  section_a4.project_purpose,
                  format_type=format_type,
                  size=11)


def write_section_a5(qapp, doc, c, format_type):
    section_a5 = SectionA5.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_a.SECTION_A['a5']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_a.SECTION_A['a5']['boilerplate'],
                  format_type=format_type,
                  size=11)
    add_paragraph(doc,
                  c,
                  section_a5.tasks_summary,
                  format_type=format_type,
                  size=11)

    start_fy = int(section_a5.start_fy)
    start_q = int(section_a5.start_q)
    quarters = ['Q1', 'Q2', 'Q3', 'Q4']
    table_data = [
        ['Project Task Schedule'] +
        [f'FY{parse_fy(start_fy) + (i - 1) // 4}' for i in range(1, 10)],
        ['Task Description'] +
        [quarters[(start_q - 1 + (i - 1)) % 4] for i in range(1, 10)]
    ]
    add_table(doc,
              c,
              table_data,
              format_type=format_type,
              column_widths=[3 * inch] + [3 / 9 * inch] * 9)


def write_section_a6(qapp, doc, c, format_type):
    section_a6 = SectionA6.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_a.SECTION_A['a6']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  section_a6.information,
                  format_type=format_type,
                  size=11)


def write_section_a7(qapp, doc, c, format_type):
    add_heading(doc,
                c,
                constants_a.SECTION_A['a7']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_a.SECTION_A['a7']['boilerplate'],
                  format_type=format_type,
                  size=11)

    distribution_list = Distribution.objects.filter(qapp_id=qapp.id)
    table_data = [[
        'Name & Organization', 'Contact Information (e-mail)', 'Project Role(s)'
    ]] + [[
        f'{recipient.name}\n{recipient.org}', recipient.email,
        recipient.proj_role
    ] for recipient in distribution_list]
    add_table(doc,
              c,
              table_data,
              format_type=format_type,
              column_widths=[2 * inch, 2 * inch, 2 * inch])


def write_section_a8(qapp, doc, c, format_type):
    add_heading(doc,
                c,
                constants_a.SECTION_A['a8']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_a.SECTION_A['a8']['boilerplate'],
                  format_type=format_type,
                  size=11)

    roles = RoleResponsibility.objects.filter(qapp_id=qapp.id)
    table_data = [[
        'Name & Organization', 'Project Role(s)', 'Project Responsibilities'
    ]] + [[
        f'{role.name}\n{role.org}', role.proj_role, '\n'.join(
            [f'• {resp}' for resp in role.proj_responsibilities.split('\n')])
    ] for role in roles]
    add_table(doc,
              c,
              table_data,
              format_type=format_type,
              column_widths=[1 * inch, 1 * inch, 4 * inch])


def write_section_a9(qapp, doc, c, format_type):
    add_heading(doc,
                c,
                constants_a.SECTION_A['a9']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_a.SECTION_A['a9']['boilerplate'],
                  format_type=format_type,
                  size=11)


def write_section_a10(qapp, doc, c, format_type):
    section_a10 = SectionA10.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_a.SECTION_A['a10']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_a.SECTION_A['a10']['boilerplate'],
                  format_type=format_type,
                  size=11)

    if format_type == 'docx' and section_a10.org_chart and exists(
            section_a10.org_chart):
        doc.add_picture(section_a10.org_chart)
    elif format_type == 'pdf' and section_a10.org_chart and exists(
            section_a10.org_chart):
        c.drawImage(section_a10.org_chart,
                    100,
                    c._height - 300,
                    width=400,
                    height=200)


def write_section_a11(qapp, doc, c, format_type):
    section_a11 = SectionA11.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_a.SECTION_A['a11']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  section_a11.information,
                  format_type=format_type,
                  size=11)


def write_section_a12(qapp, doc, c, format_type):
    add_heading(doc,
                c,
                constants_a.SECTION_A['a12']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_a.SECTION_A['a12']['boilerplate'],
                  format_type=format_type,
                  size=11)

    records = DocumentRecord.objects.filter(qapp_id=qapp.id)
    table_data = [[
        'Record Type', 'Responsible Party',
        'Located in Project File (Y/N) If No, Add File Location Below',
        'File Type (Format)', 'Special Handling Required (Y/N)'
    ]] + [[
        record.record_type, record.responsible_party, record.in_proj_file,
        record.file_type, 'Y' if record.special_handling else 'N'
    ] for record in records]
    add_table(doc,
              c,
              table_data,
              format_type=format_type,
              column_widths=[2 * inch, 2 * inch, 2 * inch, 2 * inch, 2 * inch])


def write_section_b(qapp, doc, c, format_type):
    section_b = SectionB.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_b.SECTION_B['b']['header'],
                level=1,
                format_type=format_type)

    sections_to_write = ['b1', 'b2', 'b3', 'b4', 'b5', 'b6']
    for section in sections_to_write:
        add_heading(doc,
                    c,
                    constants_b.SECTION_B[section]['header'],
                    level=2,
                    format_type=format_type)
        add_paragraph(doc,
                      c,
                      getattr(section_b, section),
                      format_type=format_type,
                      size=11)

    write_section_b7(qapp, doc, c, format_type)


def write_section_b7(qapp, doc, c, format_type):
    section_b7 = SectionB7.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_b.SECTION_B['b7']['header'],
                level=2,
                format_type=format_type)

    add_heading(doc,
                c,
                constants_b.SECTION_B['b71']['header'],
                level=3,
                format_type=format_type)
    add_paragraph(doc, c, section_b7.b71, format_type=format_type, size=11)

    add_heading(doc,
                c,
                constants_b.SECTION_B['b72']['header'],
                level=3,
                format_type=format_type)
    add_paragraph(doc, c, section_b7.b72, format_type=format_type, size=11)

    add_heading(doc,
                c,
                constants_b.SECTION_B['b73']['header'],
                level=3,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_b.SECTION_B['b73']['boilerplate'],
                  format_type=format_type,
                  size=11)

    hdw_sfw = HardwareSoftware.objects.filter(qapp_id=qapp.id)
    table_data = [[
        'Hardware', 'Operating System',
        'Non-Microsoft Office Software and Version/Special Performance Requirements/Use'  # noqa: E501
    ]] + [[row.hardware, row.os, row.details] for row in hdw_sfw]
    add_table(doc,
              c,
              table_data,
              format_type=format_type,
              column_widths=[2 * inch, 2 * inch, 4 * inch])

    add_heading(doc,
                c,
                constants_b.SECTION_B['b74']['header'],
                level=3,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_b.SECTION_B['b74']['boilerplate'],
                  format_type=format_type,
                  size=11)


def write_section_c(qapp, doc, c, format_type):
    section_c = SectionC.objects.get(qapp_id=qapp.id)
    section_a1 = SectionA1.objects.get(qapp_id=qapp.id)

    add_heading(doc,
                c,
                constants_c_d.SECTION_C['c']['header'],
                level=1,
                format_type=format_type)
    add_heading(doc,
                c,
                constants_c_d.SECTION_C['c1']['header'],
                level=2,
                format_type=format_type)
    add_heading(doc,
                c,
                constants_c_d.SECTION_C['c11']['header'],
                level=3,
                format_type=format_type)

    c11_boilerplate = constants_c_d.SECTION_C['c11']['boilerplate_a']
    if section_a1.qa_category == constants_a.QA_CATEGORY_B:
        c11_boilerplate = constants_c_d.SECTION_C['c11']['boilerplate_b']
    add_paragraph(doc, c, c11_boilerplate, format_type=format_type, size=11)

    add_heading(doc,
                c,
                constants_c_d.SECTION_C['c12']['header'],
                level=3,
                format_type=format_type)
    add_paragraph(doc,
                  c,
                  constants_c_d.SECTION_C['c12']['boilerplate'],
                  format_type=format_type,
                  size=11)

    add_heading(doc,
                c,
                constants_c_d.SECTION_C['c2']['header'],
                level=2,
                format_type=format_type)
    add_paragraph(doc, c, section_c.c2, format_type=format_type, size=11)


def write_section_d(qapp, doc, c, format_type):
    section_d = SectionD.objects.get(qapp_id=qapp.id)
    add_heading(doc,
                c,
                constants_c_d.SECTION_D['d']['header'],
                level=1,
                format_type=format_type)

    sections_to_write = ['d1', 'd2']
    for section in sections_to_write:
        add_heading(doc,
                    c,
                    constants_c_d.SECTION_D[section]['header'],
                    level=2,
                    format_type=format_type)
        add_paragraph(doc,
                      c,
                      getattr(section_d, section),
                      format_type=format_type,
                      size=11)
